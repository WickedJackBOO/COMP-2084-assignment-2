import os
import re
import json
from shutil import copyfile
from docx import Document
import azure.functions as func

def replacePlaceholders(text, jsonD):
    matches = re.findall(r"\(\{(.*?)\}\)", text)
    for key in matches:
        if key in jsonD:
            text = text.replace("({" + key + "})", jsonD[key])
    return text

def replaceInDoc(doc, jsonD):
    for para in doc.paragraphs:
        para.text = replacePlaceholders(para.text, jsonD)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = replacePlaceholders(para.text, jsonD)
    for section in doc.sections:
        for para in section.header.paragraphs:
            para.text = replacePlaceholders(para.text, jsonD)
        for para in section.footer.paragraphs:
            para.text = replacePlaceholders(para.text, jsonD)

def clearTempFolder(tempFolder):
    os.makedirs(tempFolder, exist_ok=True)
    for file in os.listdir(tempFolder):
        filePath = os.path.join(tempFolder, file)
        if os.path.isfile(filePath):
            os.remove(filePath)

def main(req: func.HttpRequest) -> func.HttpResponse:
    try:
        data = req.get_json()
        fileName = data["file"]
        fieldValues = data["fields"]

        baseDir = os.path.dirname(os.path.abspath(__file__))
        wordFolder = os.path.join(baseDir, "..", "..", "wwwroot", "WordTemplates")
        tempFolder = os.path.join(baseDir, "..", "..", "wwwroot", "TempWords")

        clearTempFolder(tempFolder)

        inFile = os.path.join(wordFolder, fileName)
        outFile = os.path.join(tempFolder, fileName)

        copyfile(inFile, outFile)
        doc = Document(outFile)
        replaceInDoc(doc, fieldValues)
        doc.save(outFile)

        return func.HttpResponse("/TempWords/" + fileName)
    except Exception as e:
        return func.HttpResponse(str(e), status_code=500)
