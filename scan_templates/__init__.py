import os
import re
import json
from docx import Document
import azure.functions as func

def findPlaceholders(text):
    matches = re.findall(r"\(\{.*?\}\)", text)
    return [m[2:-2] for m in matches]

def extractPlaceholders(filePath):
    doc = Document(filePath)
    placeholders = set()
    for para in doc.paragraphs:
        placeholders.update(findPlaceholders(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(findPlaceholders(para.text))
    for section in doc.sections:
        for para in section.header.paragraphs:
            placeholders.update(findPlaceholders(para.text))
        for para in section.footer.paragraphs:
            placeholders.update(findPlaceholders(para.text))
    return sorted(placeholders)

def main(req: func.HttpRequest) -> func.HttpResponse:
    try:
        baseDir = os.path.dirname(os.path.abspath(__file__))
        wordFolder = os.path.join(baseDir, "..", "..", "wwwroot", "WordTemplates")
        result = {}
        for fileName in os.listdir(wordFolder):
            if fileName.endswith(".docx"):
                filePath = os.path.join(wordFolder, fileName)
                result[fileName] = extractPlaceholders(filePath)
        return func.HttpResponse(json.dumps(result), mimetype="application/json")
    except Exception as e:
        return func.HttpResponse(str(e), status_code=500)
